"""Text extraction for AK03.c uploads.

Audio/video transcription is deferred to v1.1 — the worker raises a clear
'not supported' error for those mime types so the UI can show the right
message instead of a stack trace.
"""

from __future__ import annotations

import io
import logging
from typing import Final

logger = logging.getLogger(__name__)

# Cap the text we feed Claude so a 200-page PDF doesn't blow our token budget.
# Claude Sonnet 4.5 has plenty of context, but we only summarise meeting docs;
# the first ~24k characters cover everything that matters.
MAX_EXTRACT_CHARS: Final[int] = 24_000


class ExtractError(Exception):
    """Raised when we can't get text out — surfaced to the user verbatim."""


def extract_text(filename: str, mime_type: str | None, data: bytes) -> str:
    name = (filename or "").lower()

    # Defer audio/video to v1.1 explicitly.
    if name.endswith((".mp3", ".mp4", ".m4a", ".wav", ".mov")) or (
        mime_type and (mime_type.startswith("audio/") or mime_type.startswith("video/"))
    ):
        raise ExtractError(
            "Audio/video transcription lands in v1.1. Upload a .docx/.pdf/.txt/.vtt instead."
        )

    if name.endswith(".docx"):
        text = _extract_docx(data)
    elif name.endswith(".doc"):
        # Legacy binary Word — no reliable pure-Python parser; ask user to convert.
        raise ExtractError(
            "Legacy .doc format isn't supported for text extraction. "
            "Open the file in Word and 'Save As' .docx, then re-upload."
        )
    elif name.endswith(".pdf"):
        text = _extract_pdf(data)
    elif name.endswith(".vtt"):
        text = _extract_vtt(data)
    elif name.endswith(".eml"):
        text = _extract_eml(data)
    elif name.endswith(".txt"):
        text = data.decode("utf-8", errors="replace")
    else:
        raise ExtractError(f"Unsupported extension for text extraction: {filename}")

    text = (text or "").strip()
    if not text:
        raise ExtractError("File contained no extractable text.")
    if len(text) > MAX_EXTRACT_CHARS:
        text = text[:MAX_EXTRACT_CHARS] + "\n\n…[truncated]"
    return text


def _extract_docx(data: bytes) -> str:
    from docx import Document  # python-docx

    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    # Tables are common in MOMs; append cell text in row order.
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages: list[str] = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            # Some PDFs contain only images; we accept partial extraction.
            logger.warning("PDF page extraction failed; skipping a page")
            continue
    return "\n\n".join(pages)


def _extract_vtt(data: bytes) -> str:
    """Strip WebVTT cue headers + timestamps; keep only the spoken text lines."""
    raw = data.decode("utf-8", errors="replace")
    out: list[str] = []
    for line in raw.splitlines():
        s = line.strip()
        if not s or s == "WEBVTT" or "-->" in s or s.isdigit():
            continue
        out.append(s)
    return " ".join(out)


def _extract_eml(data: bytes) -> str:
    """Parse an Outlook/RFC-5322 .eml file. Prefer text/plain; fall back to a
    naive HTML strip for text/html. Prepend a small header summary so the AI
    sees subject + participants alongside the body."""
    import re
    from email import message_from_bytes
    from email.policy import default as default_policy

    msg = message_from_bytes(data, policy=default_policy)

    header_lines: list[str] = []
    for label, key in (("From", "From"), ("To", "To"), ("Cc", "Cc"), ("Subject", "Subject"), ("Date", "Date")):
        val = msg.get(key)
        if val:
            header_lines.append(f"{label}: {val}")

    # Walk parts, prefer text/plain. Capture text/html as fallback.
    plain_parts: list[str] = []
    html_parts: list[str] = []
    for part in msg.walk():
        if part.is_multipart():
            continue
        ctype = (part.get_content_type() or "").lower()
        # Skip attachments by Content-Disposition.
        disp = (part.get("Content-Disposition") or "").lower()
        if "attachment" in disp:
            continue
        try:
            content = part.get_content()
        except Exception:
            payload = part.get_payload(decode=True) or b""
            content = payload.decode("utf-8", errors="replace") if isinstance(payload, bytes) else str(payload)
        if not isinstance(content, str):
            continue
        if ctype == "text/plain":
            plain_parts.append(content)
        elif ctype == "text/html":
            html_parts.append(content)

    body = "\n\n".join(plain_parts).strip()
    if not body and html_parts:
        # Naive HTML → text: drop tags + collapse whitespace. Good enough for MOMs.
        raw_html = "\n\n".join(html_parts)
        # Drop <style>/<script> blocks entirely.
        raw_html = re.sub(r"(?is)<(style|script)[^>]*>.*?</\1>", " ", raw_html)
        # Convert <br> and </p> to newlines so paragraph structure survives.
        raw_html = re.sub(r"(?i)<br\s*/?>", "\n", raw_html)
        raw_html = re.sub(r"(?i)</p\s*>", "\n\n", raw_html)
        # Strip remaining tags.
        text_only = re.sub(r"<[^>]+>", " ", raw_html)
        # Unescape common entities.
        text_only = (
            text_only.replace("&nbsp;", " ")
            .replace("&amp;", "&")
            .replace("&lt;", "<")
            .replace("&gt;", ">")
            .replace("&quot;", '"')
            .replace("&#39;", "'")
        )
        # Collapse runs of whitespace per line; keep paragraph breaks.
        lines = [re.sub(r"\s+", " ", ln).strip() for ln in text_only.splitlines()]
        body = "\n".join(ln for ln in lines if ln)

    header_block = "=== HEADERS ===\n" + "\n".join(header_lines) if header_lines else ""
    body_block = f"\n\n=== BODY ===\n{body}" if body else ""
    return (header_block + body_block).strip()
